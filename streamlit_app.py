import streamlit as st
import pandas as pd
import io
import os
import re
import base64
import pathlib
import unicodedata
import zipfile
from datetime import datetime, timedelta, date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# ── Paleta de marca (logo Smart HSE) ─────────────────────────
BRAND = {"navy": "#0E3A5F", "blue": "#16609E", "cyan": "#27AAE1", "green": "#5BBA47"}

# ── Assets de marca ──────────────────────────────────────────
_BASE = os.path.dirname(os.path.abspath(__file__))
def _b64(rel):
    try:
        return base64.b64encode((pathlib.Path(_BASE) / rel).read_bytes()).decode()
    except Exception:
        return ""

def _uri(rel):
    b = _b64(rel)
    return f"data:image/png;base64,{b}" if b else ""
_FULL = _uri("assets/logo_smarthse.png")
_MARK = _uri("assets/logo_mark.png")
LOGO_FULL_URI = _FULL or _MARK
LOGO_MARK_URI = _MARK or _FULL
_HAS_FULL_LOGO = bool(_FULL)
_FAVICON = os.path.join(_BASE, "assets", "favicon.png")

st.set_page_config(
    page_title="Smart HSE Chile — Gestión HSE para todas las áreas laborales",
    page_icon=_FAVICON if os.path.exists(_FAVICON) else "🛡️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

def logo_img(height=46, mark=False):
    uri = LOGO_MARK_URI if mark else LOGO_FULL_URI
    return f"<img src='{uri}' style='height:{height}px;width:auto;display:inline-block'/>" if uri else ""

# ── Estado ──────────────────────────────────────────────────
def init():
    for k,v in {"vista":"landing","auth":False,"contrato":"","actividad":"","lugar":"","ruta":"","incidentes":[],"leads":[]}.items():
        if k not in st.session_state: st.session_state[k]=v
init()

APP_PW = os.environ.get("APP_PASSWORD", "smarthse2025")

# ── CSS ─────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@700;800;900&family=Inter:wght@300;400;500;600;700&display=swap');
:root{--navy:#0E3A5F;--blue:#16609E;--cyan:#27AAE1;--green:#5BBA47;--ink:#0E3A5F;--muted:#5b7184;--line:#e2e8f0;}
#MainMenu,footer,header{visibility:hidden}
.block-container{padding:0!important;max-width:100%!important}
section[data-testid="stSidebar"]{display:none}
html,body,[class*="css"]{font-family:'Inter',sans-serif}
a{text-decoration:none}
@keyframes fadeUp{from{opacity:0;transform:translateY(26px)}to{opacity:1;transform:translateY(0)}}
.fu{animation:fadeUp .7s cubic-bezier(.2,.7,.2,1) both}
.fu2{animation:fadeUp .7s .12s cubic-bezier(.2,.7,.2,1) both}
.fu3{animation:fadeUp .7s .24s cubic-bezier(.2,.7,.2,1) both}

/* ── Nav ── */
.nav{display:flex;align-items:center;justify-content:space-between;padding:14px 46px;background:#fff;box-shadow:0 2px 18px rgba(14,58,95,.07);position:sticky;top:0;z-index:999;flex-wrap:wrap;gap:12px}
.nav-logo{display:flex;align-items:center;gap:11px}
.wm{font-family:'Montserrat',sans-serif;font-weight:900;letter-spacing:.5px;line-height:1}
.wm .smart{color:var(--blue)} .wm .hse{color:var(--cyan)}
.chip-chile{background:var(--cyan);color:#fff;font-weight:700;border-radius:4px;letter-spacing:3px;display:inline-block}
.nav-links{display:flex;gap:26px;align-items:center;flex-wrap:wrap}
.nav-links a{color:#42566a;font-size:12.5px;font-weight:700;text-transform:uppercase;letter-spacing:.6px;transition:color .2s}
.nav-links a:hover,.nav-links a.active{color:var(--cyan)}
.btn-demo{background:var(--green);color:#fff!important;padding:11px 24px;border-radius:30px;font-size:12.5px;font-weight:800;text-transform:uppercase;letter-spacing:.8px;box-shadow:0 8px 20px rgba(91,186,71,.35);transition:transform .2s,box-shadow .2s}
.btn-demo:hover{transform:translateY(-2px);box-shadow:0 12px 26px rgba(91,186,71,.45)}

/* ── Hero ── */
.hero{position:relative;overflow:hidden;background:linear-gradient(125deg,var(--navy) 0%,var(--blue) 55%,#1f7fc0 100%);padding:104px 20px 150px;text-align:center;color:#fff}
.hero::before{content:"";position:absolute;inset:0;background-image:linear-gradient(rgba(14,58,95,.82),rgba(22,96,158,.78)),url('https://images.unsplash.com/photo-1504917595217-d4dc5ebe6122?auto=format&fit=crop&w=2000&q=80');background-size:cover;background-position:center;opacity:.5}
.hero::after{content:"";position:absolute;top:-28%;right:-10%;width:520px;height:520px;background:radial-gradient(circle,rgba(91,186,71,.35),transparent 60%);pointer-events:none}
.hero>*{position:relative;z-index:2}
.hero-logo{width:122px;height:122px;margin:0 auto 22px;border-radius:50%;background:rgba(255,255,255,.96);display:flex;align-items:center;justify-content:center;box-shadow:0 16px 44px rgba(0,0,0,.30);border:1px solid rgba(255,255,255,.5)}
.eyebrow{display:inline-block;background:rgba(255,255,255,.12);border:1px solid rgba(255,255,255,.28);color:#dff3ff;font-size:11px;font-weight:600;letter-spacing:2px;text-transform:uppercase;padding:7px 16px;border-radius:30px;margin-bottom:22px;backdrop-filter:blur(6px)}
.hero h1{font-family:'Montserrat',sans-serif;font-weight:900;font-size:50px;max-width:960px;margin:0 auto 22px;text-transform:uppercase;line-height:1.12;text-shadow:0 4px 18px rgba(0,0,0,.3)}
.hero h1 .hl{color:var(--green)}
.hero p{font-size:18px;max-width:730px;margin:0 auto 38px;font-weight:300;line-height:1.7;color:#eaf4fb}
.btn-hero{background:var(--green);color:#fff;padding:16px 38px;border-radius:30px;font-weight:800;font-size:14px;text-transform:uppercase;letter-spacing:1px;display:inline-block;box-shadow:0 12px 28px rgba(91,186,71,.45);transition:transform .25s,box-shadow .25s;margin:0 8px}
.btn-hero:hover{transform:translateY(-3px);box-shadow:0 18px 36px rgba(91,186,71,.55)}
.btn-ghost{background:transparent;color:#fff;padding:14px 34px;border-radius:30px;font-weight:700;font-size:14px;text-transform:uppercase;letter-spacing:1px;display:inline-block;border:2px solid rgba(255,255,255,.45);transition:background .25s;margin:0 8px}
.btn-ghost:hover{background:rgba(255,255,255,.14)}
.claims{display:flex;justify-content:center;flex-wrap:wrap;gap:13px;max-width:920px;margin:38px auto 0}
.claim{display:flex;align-items:center;gap:9px;background:rgba(255,255,255,.10);border:1px solid rgba(255,255,255,.22);color:#eaf6ff;font-size:13px;font-weight:600;padding:10px 18px;border-radius:30px;backdrop-filter:blur(6px)}
.claim .dot{width:8px;height:8px;border-radius:50%;background:var(--green);box-shadow:0 0 10px var(--green)}

/* ── Secciones ── */
.sec{max-width:1140px;margin:0 auto;padding:64px 22px 12px}
.sec-tag{text-align:center;color:var(--cyan);font-weight:800;font-size:12px;letter-spacing:3px;text-transform:uppercase;margin-bottom:8px}
.sec-h{text-align:center;font-family:'Montserrat',sans-serif;font-weight:800;font-size:32px;color:var(--ink);margin:0 0 10px}
.sec-sub{text-align:center;color:var(--muted);font-size:15.5px;max-width:660px;margin:0 auto 42px;line-height:1.6}

/* ── Sectores ── */
.sectores{display:grid;grid-template-columns:repeat(4,1fr);gap:16px}
.sector{background:#fff;border:1px solid var(--line);border-radius:14px;padding:24px 16px;text-align:center;transition:transform .28s,box-shadow .28s,border-color .28s}
.sector:hover{transform:translateY(-6px);box-shadow:0 18px 38px rgba(14,58,95,.12);border-color:var(--cyan)}
.sector .si{font-size:30px;display:block;margin-bottom:10px}
.sector b{display:block;font-size:13.5px;color:var(--ink);font-weight:700}

/* ── Cards servicios ── */
.cards{display:flex;justify-content:center;gap:22px;flex-wrap:wrap}
.card{background:#fff;padding:36px 24px;border-radius:18px;width:23%;min-width:230px;text-align:center;box-shadow:0 14px 40px rgba(14,58,95,.10);border:1px solid #eef3f8;border-top:4px solid transparent;transition:transform .3s,border-color .3s,box-shadow .3s}
.card:hover{transform:translateY(-8px);border-top-color:var(--cyan);box-shadow:0 26px 56px rgba(22,96,158,.20)}
.card-icon{font-size:28px;width:62px;height:62px;line-height:62px;margin:0 auto 6px;border-radius:16px;background:linear-gradient(135deg,var(--cyan),var(--blue));color:#fff;box-shadow:0 8px 20px rgba(39,170,225,.35)}
.card h3{font-family:'Montserrat',sans-serif;font-weight:800;font-size:14px;color:var(--ink);text-transform:uppercase;margin:16px 0 8px}
.card p{font-size:13px;color:var(--muted);line-height:1.6}

/* ── Cómo funciona ── */
.steps{display:flex;gap:22px;flex-wrap:wrap;justify-content:center}
.step{flex:1;min-width:240px;background:#fff;border:1px solid var(--line);border-radius:16px;padding:30px 24px;position:relative;transition:transform .3s,box-shadow .3s}
.step:hover{transform:translateY(-6px);box-shadow:0 18px 40px rgba(14,58,95,.10)}
.step .num{font-family:'Montserrat',sans-serif;font-weight:900;font-size:42px;line-height:1;background:linear-gradient(135deg,var(--cyan),var(--green));-webkit-background-clip:text;background-clip:text;-webkit-text-fill-color:transparent;margin-bottom:12px}
.step h4{font-family:'Montserrat',sans-serif;font-weight:800;font-size:16px;color:var(--ink);margin:0 0 8px}
.step p{font-size:13.5px;color:var(--muted);line-height:1.6;margin:0}

/* ── Beneficios ── */
.why{background:linear-gradient(120deg,var(--navy),var(--blue));border-radius:24px;max-width:1140px;margin:64px auto 0;padding:48px 40px;color:#fff;display:grid;grid-template-columns:repeat(3,1fr);gap:30px}
.why .b .bi{font-size:26px;margin-bottom:8px}
.why .b h4{font-family:'Montserrat',sans-serif;font-size:16px;margin:0 0 6px;font-weight:800}
.why .b p{font-size:13.5px;color:#cfe4f3;line-height:1.6;margin:0}

/* ── Footer ── */
.sh-footer{background:var(--navy);color:#9fb6c9;text-align:center;padding:54px 20px;font-size:13px;margin-top:64px}
.sh-footer a{color:var(--cyan)}
.ftr-sep{border-top:1px solid rgba(255,255,255,.10);margin-top:22px;padding-top:16px;font-size:11px;color:#6f879b}

/* ── Botones nativos Streamlit (form) → verde ── */
.stForm button[kind="primaryFormSubmit"]{background:var(--green)!important;border-color:var(--green)!important}
.stForm button[kind="primaryFormSubmit"]:hover{background:#4ea23b!important;border-color:#4ea23b!important}
.stDownloadButton>button{background:var(--cyan)!important;border-color:var(--cyan)!important}

@media(max-width:900px){.sectores{grid-template-columns:repeat(2,1fr)}.why{grid-template-columns:1fr}}
@media(max-width:768px){
  .nav{padding:12px 18px;justify-content:center}.nav-links{gap:14px;justify-content:center}
  .card{width:100%}.cards{flex-direction:column}
  .hero h1{font-size:30px}.hero{padding:80px 18px 120px}
  .btn-hero,.btn-ghost{display:block;margin:8px auto;max-width:320px}
  .steps{flex-direction:column}.sec-h{font-size:25px}
}

/* ── Consola ── */
.kpi{background:linear-gradient(135deg,#0E3A5F 0%,#16609E 100%);border-radius:12px;padding:1.2rem 1.5rem;color:white;text-align:center;box-shadow:0 4px 14px rgba(14,58,95,.15);margin-bottom:8px}
.kpi .val{font-size:2.4rem;font-weight:300}
.kpi .lbl{font-size:.8rem;opacity:.8;margin-top:4px}
.ar{background:#fef2f2;border-left:4px solid #ef4444;padding:.8rem 1rem;border-radius:6px;margin:.4rem 0}
.ag{background:#f0fdf4;border-left:4px solid #5BBA47;padding:.8rem 1rem;border-radius:6px;margin:.4rem 0}
.aa{background:#fffbeb;border-left:4px solid #f59e0b;padding:.8rem 1rem;border-radius:6px;margin:.4rem 0}
.nota{background:#fff8e1;border-left:4px solid #FFA000;padding:12px 16px;border-radius:6px;font-size:.85rem}
</style>
""", unsafe_allow_html=True)

# ── Datos maestros ───────────────────────────────────────────
ESTRUCTURA = {
    "Minera Spence":  {"faenas":["Spence BHP"],"ruta":"Spence/Carpeta_Arranque"},
    "Codelco DRT":    {"faenas":["Radomiro Tomic","Chuquicamata","Ministro Hales"],"ruta":"Codelco/Contratos_Estandar"},
    "Minera El Abra": {"faenas":["Mina","Planta"],"ruta":"El_Abra/Documentos"},
    "Centinela":      {"faenas":["Mina","Planta SX-EW"],"ruta":"Centinela/Documentos"},
}
REQUISITOS = {
    "Minera Spence":  ["Validación fotográfica ECF21","Firma administrador en matriz base","Acreditación vigente de equipos"],
    "Codelco DRT":    ["LOD semanal enviado","Reporte de observaciones conductuales","Matriz de riesgos actualizada"],
    "Minera El Abra": ["Charla 5 minutos firmada","Validación ingreso a planta"],
    "Centinela":      ["Registro All Scan diario","Revisión EPP específico"],
}
ACTIVIDADES = [
    {"actividad":"Registro FYS / ECF21 diario","frecuencia":"Diaria","dias":1,"modulo":"FYS","cliente":"Codelco DRT"},
    {"actividad":"Actualización ECF21 Terreno","frecuencia":"Diaria","dias":1,"modulo":"Terreno","cliente":"Minera Spence"},
    {"actividad":"Revisión LOD semanal","frecuencia":"Semanal","dias":7,"modulo":"RESSO","cliente":"Codelco DRT"},
    {"actividad":"Acreditación Equipos","frecuencia":"Semanal","dias":5,"modulo":"Legal","cliente":"Minera Spence"},
    {"actividad":"Reunión CPHS","frecuencia":"Mensual","dias":30,"modulo":"Legal","cliente":"Minera El Abra"},
    {"actividad":"Registro All Scan","frecuencia":"Diaria","dias":1,"modulo":"Terreno","cliente":"Centinela"},
    {"actividad":"Revisión EPP específico","frecuencia":"Semanal","dias":5,"modulo":"Legal","cliente":"Centinela"},
]
CONTRATOS = ["405","118","109100077748"]
FLUJOS    = ["FYS Diario","RESSO"]

def fecha_es(d=None):
    d = d or date.today()
    meses={"January":"enero","February":"febrero","March":"marzo","April":"abril","May":"mayo","June":"junio","July":"julio","August":"agosto","September":"septiembre","October":"octubre","November":"noviembre","December":"diciembre"}
    t=d.strftime("%d de %B de %Y")
    for en,es in meses.items(): t=t.replace(en,es)
    return t

# ── Generador GENERAL de Cartas de No Aplicabilidad ──────────
# Sirve para cualquier empresa y cualquier cliente/mandante; no está
# amarrado a Codelco ni a RESSO. La referencia, el fundamento y la
# declaración salen de cada fila del Excel.
CARTA_COLS = ["N°","Fecha","Señores","Ref.","Nombre del contrato","N° contrato",
              "Empresa","Declaración de No Aplica","Nombre del cliente",
              "Responsable","Cargo","Punto de RESSO"]

def _norm(s):
    s = str(s).strip().lower()
    s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
    return re.sub(r"[^a-z0-9]+", " ", s).strip()

_ALIAS = {
    "n":"numero","numero":"numero","item":"numero",
    "punto de resso":"referencia_estandar","punto resso":"referencia_estandar","estandar":"referencia_estandar",
    "fecha":"fecha",
    "senores":"senores","senores as":"senores","destinatario":"senores",
    "ref":"ref","referencia":"ref",
    "nombre del contrato":"nombre_contrato","contrato":"nombre_contrato",
    "n contrato":"num_contrato","numero contrato":"num_contrato",
    "empresa":"empresa","empresa 1":"empresa",
    "declaracion de no aplica":"declaracion","declaracion":"declaracion","no aplica":"declaracion",
    "nombre delcliente":"nombre_cliente","nombre del cliente":"nombre_cliente","cliente":"nombre_cliente",
    "responsable":"responsable",
    "carog":"cargo","cargo":"cargo",
}

def mapear_columnas(df):
    nuevo = {}
    for col in df.columns:
        clave = _ALIAS.get(_norm(col))
        if clave and clave not in nuevo.values():
            nuevo[col] = clave
    return df.rename(columns=nuevo)

def _cval(row, clave, default=""):
    v = row.get(clave, default)
    if v is None or (isinstance(v, float) and pd.isna(v)):
        return default
    return str(v).strip()

def _limpiar(t):
    t = re.sub(r"\s+", " ", str(t)).strip()
    t = re.sub(r"\.{2,}", ".", t)
    t = re.sub(r"\s+([.,;:])", r"\1", t)
    return t

def _slug(s):
    return re.sub(r"_+", "_", _norm(s).replace(" ", "_")).strip("_") or "empresa"

def construir_carta_doc(row, ciudad="Antofagasta"):
    """Devuelve un Document de python-docx con la carta formal de la fila."""
    doc = Document()
    base = doc.styles["Normal"]; base.font.name = "Calibri"; base.font.size = Pt(11)

    empresa = _cval(row, "empresa", "La Empresa")
    fecha = _cval(row, "fecha") or fecha_es(date.today())
    senores = _cval(row, "senores") or _cval(row, "nombre_cliente", "Señores(as)")
    ref = _cval(row, "ref")
    referencia_estandar = _cval(row, "referencia_estandar")
    nombre_contrato = _cval(row, "nombre_contrato")
    num_contrato = re.sub(r"^\s*n[°ºo\.]*\s*", "", _cval(row, "num_contrato"), flags=re.I)
    declaracion = _cval(row, "declaracion", "no aplica según la naturaleza del servicio.")
    declaracion = re.sub(r"^\s*(informa[,\s]+que|informa)[\s,:]*", "", declaracion, flags=re.I)
    if declaracion:
        declaracion = declaracion[0].lower() + declaracion[1:]
    cliente = _cval(row, "nombre_cliente")
    responsable = _cval(row, "responsable")
    cargo = _cval(row, "cargo", "Asesor en Prevención de Riesgos")

    p = doc.add_paragraph(f"{ciudad}, {fecha}"); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("Señores(as)")
    doc.add_paragraph(senores)
    doc.add_paragraph("Presente")
    doc.add_paragraph("")
    partes = [x for x in (ref, f"Ítem {referencia_estandar}" if referencia_estandar else "") if x]
    pref = doc.add_paragraph()
    pref.add_run("Ref.: " + (" — ".join(partes) if partes else "Declaración de No Aplicabilidad")).bold = True
    doc.add_paragraph("")
    doc.add_paragraph("De nuestra consideración:")
    doc.add_paragraph("")
    cuerpo = f"Mediante la presente, {empresa} "
    if nombre_contrato:
        cuerpo += f"en el marco del contrato “{nombre_contrato}”"
        if num_contrato: cuerpo += f" N° {num_contrato}"
        cuerpo += ", "
    cuerpo += f"informa que {declaracion.rstrip('. ')}"
    if cliente and cliente.lower().rstrip(".") not in cuerpo.lower():
        cuerpo += f" ante {cliente}"
    if not cuerpo.rstrip().endswith("."): cuerpo += "."
    doc.add_paragraph(_limpiar(cuerpo))
    doc.add_paragraph("")
    doc.add_paragraph("Sin otro particular, le saluda cordialmente,")
    for _ in range(3): doc.add_paragraph("")
    doc.add_paragraph("________________________")
    if responsable: doc.add_paragraph(responsable)
    doc.add_paragraph(cargo)
    doc.add_paragraph(empresa)
    return doc

def nombre_archivo_carta(row, idx):
    empresa = _slug(_cval(row, "empresa", "empresa"))
    num = _cval(row, "numero", str(idx + 1)).replace(".0", "").replace(".", "")
    fecha = _slug(_cval(row, "fecha", date.today().isoformat()))
    return f"Carta_No_Aplica_{empresa}_{num}_{fecha}.docx"

def plantilla_excel_bytes():
    """Excel de ejemplo con las columnas esperadas, para descargar como guía."""
    ejemplo = pd.DataFrame([{
        "N°":1,"Fecha":fecha_es(date.today()),"Señores":"Nombre del Cliente / Mandante",
        "Ref.":"Motivo de la no aplicabilidad","Nombre del contrato":"Nombre del servicio",
        "N° contrato":"0000","Empresa":"Tu Empresa SPA",
        "Declaración de No Aplica":"no aplica el ítem por la naturaleza del servicio.",
        "Nombre del cliente":"Nombre del Cliente / Mandante","Responsable":"Nombre Apellido",
        "Cargo":"Asesor en Prevención de Riesgos","Punto de RESSO":"(opcional)"}])
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        ejemplo.to_excel(w, index=False, sheet_name="Carta de N A")
    return buf.getvalue()

# ════════════════════════════════════════════════════════════
# VISTA LANDING (marketing transversal + acceso a consola)
# ════════════════════════════════════════════════════════════
def landing():
    _cta_l, _cta_r = st.columns([6,1])
    with _cta_r:
        if st.button("🔒 Consola", use_container_width=True):
            st.session_state["vista"]="login"; st.rerun()

    # ════════════════════════════════════════════════════════════
    # NAV
    # ════════════════════════════════════════════════════════════
    _logo_block = logo_img(52) if _HAS_FULL_LOGO else (
        f"{logo_img(46, mark=True)}<div class='wm'><span style='font-size:21px'><span class='smart'>SMART</span> <span class='hse'>HSE</span></span>"
        f"<br><span class='chip-chile' style='font-size:9px;padding:1px 8px'>CHILE</span></div>"
    )
    st.markdown(f"""
    <div class="nav">
      <div class="nav-logo">{_logo_block}</div>
      <div class="nav-links">
        <a href="#inicio" class="active">Inicio</a>
        <a href="#soluciones">Soluciones</a>
        <a href="#tecnologia">Tecnología</a>
        <a href="#nosotros">Nosotros</a>
        <a href="#sectores">Sectores</a>
        <a href="#contacto">Contacto</a>
        <a href="#contacto" class="btn-demo">Solicitar demo</a>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════
    # HERO
    # ════════════════════════════════════════════════════════════
    st.markdown(f"""
    <div id="inicio"></div>
    <div class="hero">
      <div class="hero-logo fu">{logo_img(80, mark=True)}</div>
      <div class="eyebrow fu">🛡️ Seguridad · Salud Ocupacional · Medio Ambiente</div>
      <h1 class="fu">Gestión HSE inteligente para <span class="hl">todas las áreas laborales</span> de Chile</h1>
      <p class="fu2">Acompañamos a empresas de cualquier sector a cumplir el DS.44, prevenir riesgos y construir cultura de seguridad — con tecnología, trazabilidad total y asesoría experta.</p>
      <div class="fu3">
        <a href="#contacto" class="btn-hero">Solicitar demo</a>
        <a href="#soluciones" class="btn-ghost">Conocer soluciones</a>
      </div>
      <div class="claims fu3">
        <div class="claim"><span class="dot"></span>Cumplimiento DS.44</div>
        <div class="claim"><span class="dot"></span>Multisector · Transversal</div>
        <div class="claim"><span class="dot"></span>Trazabilidad documental total</div>
        <div class="claim"><span class="dot"></span>Asesoría experta en terreno</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════
    # SECTORES
    # ════════════════════════════════════════════════════════════
    st.markdown("""
    <div id="sectores"></div>
    <div class="sec">
      <div class="sec-tag">Transversal</div>
      <div class="sec-h">Una solución para todos los sectores</div>
      <div class="sec-sub">El DS.44 aplica a toda empresa con trabajadores. Sin importar tu rubro, Smart HSE adapta la gestión de seguridad y salud ocupacional a tu realidad.</div>
      <div class="sectores">
        <div class="sector"><span class="si">⛏️</span><b>Minería</b></div>
        <div class="sector"><span class="si">🏗️</span><b>Construcción</b></div>
        <div class="sector"><span class="si">🏭</span><b>Industria y Manufactura</b></div>
        <div class="sector"><span class="si">🚚</span><b>Logística y Transporte</b></div>
        <div class="sector"><span class="si">⚡</span><b>Energía</b></div>
        <div class="sector"><span class="si">🌾</span><b>Agro y Alimentos</b></div>
        <div class="sector"><span class="si">🏢</span><b>Servicios y Oficinas</b></div>
        <div class="sector"><span class="si">🛒</span><b>Retail y Comercio</b></div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════
    # SOLUCIONES / SERVICIOS
    # ════════════════════════════════════════════════════════════
    st.markdown("""
    <div id="soluciones"></div>
    <div class="sec">
      <div class="sec-tag">Soluciones</div>
      <div class="sec-h">Lo que hacemos por tu empresa</div>
      <div class="sec-sub">Un sistema completo de gestión HSE que cubre desde la identificación de riesgos hasta la cultura preventiva.</div>
      <div class="cards">
        <div class="card"><div class="card-icon">⚠️</div><h3>Gestión de Riesgos</h3><p>Identificación, evaluación y control de peligros con matrices y planes de acción según DS.44.</p></div>
        <div class="card"><div class="card-icon">📋</div><h3>Cumplimiento Normativo</h3><p>Seguimiento en tiempo real de obligaciones legales, plazos y vencimientos críticos.</p></div>
        <div class="card"><div class="card-icon">📊</div><h3>Análisis y Datos</h3><p>Indicadores de seguridad, reportes ejecutivos y trazabilidad documental completa.</p></div>
        <div class="card"><div class="card-icon">🛡️</div><h3>Cultura de Seguridad</h3><p>Capacitación, observaciones conductuales y gestión del comportamiento seguro.</p></div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════
    # CÓMO FUNCIONA / TECNOLOGÍA
    # ════════════════════════════════════════════════════════════
    st.markdown("""
    <div id="tecnologia"></div>
    <div class="sec">
      <div class="sec-tag">Cómo funciona</div>
      <div class="sec-h">Tecnología que ordena tu gestión HSE</div>
      <div class="sec-sub">Un proceso simple y acompañado, de principio a fin.</div>
      <div class="steps">
        <div class="step"><div class="num">01</div><h4>Diagnóstico</h4><p>Evaluamos tu cumplimiento DS.44 actual, detectamos brechas y priorizamos lo urgente para tu sector.</p></div>
        <div class="step"><div class="num">02</div><h4>Implementación</h4><p>Centralizamos documentos, matrices y registros en la plataforma, con clasificación y trazabilidad automática.</p></div>
        <div class="step"><div class="num">03</div><h4>Mejora continua</h4><p>Monitoreas indicadores, generas reportes y mantienes la cultura preventiva viva en el tiempo.</p></div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════
    # POR QUÉ / NOSOTROS
    # ════════════════════════════════════════════════════════════
    st.markdown("""
    <div id="nosotros"></div>
    <div class="sec" style="padding-bottom:0">
      <div class="sec-tag">Por qué Smart HSE</div>
      <div class="sec-h">Experiencia + tecnología, a tu lado</div>
    </div>
    <div class="why">
      <div class="b"><div class="bi">🤝</div><h4>Acompañamiento experto</h4><p>Asesores especialistas en seguridad y salud ocupacional que conocen el terreno y la normativa chilena.</p></div>
      <div class="b"><div class="bi">⚙️</div><h4>Tecnología a tu favor</h4><p>Plataforma que automatiza la documentación, reduce planillas y te da control en tiempo real.</p></div>
      <div class="b"><div class="bi">🔒</div><h4>Trazabilidad total</h4><p>Cada registro queda respaldado y disponible, listo para auditorías y fiscalizaciones.</p></div>
    </div>
    """, unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════
    # CONTACTO
    # ════════════════════════════════════════════════════════════
    st.markdown("""
    <div id="contacto"></div>
    <div class="sec">
      <div class="sec-tag">Hablemos</div>
      <div class="sec-h">Solicita una demostración</div>
      <div class="sec-sub">Cuéntanos de tu empresa y te mostramos cómo Smart HSE simplifica el cumplimiento DS.44 y la prevención en tu sector.</div>
    </div>
    """, unsafe_allow_html=True)

    SECTORES = ["Minería", "Construcción", "Industria y Manufactura", "Logística y Transporte",
                "Energía", "Agro y Alimentos", "Servicios y Oficinas", "Retail y Comercio", "Otro"]

    _, fc, _ = st.columns([1, 2, 1])
    with fc:
        with st.form("lead"):
            a, b = st.columns(2)
            nombre = a.text_input("Nombre", placeholder="Tu nombre")
            empresa = b.text_input("Empresa", placeholder="Nombre de tu empresa")
            c, d = st.columns(2)
            correo = c.text_input("Correo electrónico", placeholder="correo@empresa.cl")
            sector = d.selectbox("Sector", SECTORES)
            msg = st.text_area("¿Qué necesitas resolver?",
                               placeholder="Ej: ordenar el cumplimiento DS.44 de mi empresa y reducir el papeleo.",
                               height=90)
            enviar = st.form_submit_button("Solicitar demo →", type="primary", use_container_width=True)
        if enviar:
            if not nombre.strip() or not correo.strip():
                st.error("Por favor completa al menos tu nombre y correo.")
            else:
                st.session_state["leads"].append({
                    "nombre": nombre, "empresa": empresa, "correo": correo,
                    "sector": sector, "mensaje": msg,
                    "fecha": datetime.now().strftime("%Y-%m-%d %H:%M")})
                asunto = f"Solicitud de demo — {empresa or nombre}".replace(" ", "%20")
                cuerpo = (f"Nombre: {nombre}%0D%0AEmpresa: {empresa}%0D%0ASector: {sector}"
                          f"%0D%0ACorreo: {correo}%0D%0A%0D%0AMensaje:%0D%0A{msg}").replace(" ", "%20")
                mailto = f"mailto:contacto@smarthse.cl?subject={asunto}&body={cuerpo}"
                st.success(f"✅ ¡Gracias, {nombre.split()[0]}! Recibimos tu solicitud. Te contactaremos a la brevedad.")
                st.markdown(f"<a href='{mailto}' class='btn-hero' style='margin-top:6px'>📧 Enviar también por correo</a>",
                            unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════
    # FOOTER
    # ════════════════════════════════════════════════════════════
    st.markdown(f"""
    <div class="sh-footer">
      <div style="margin-bottom:12px">{logo_img(60, mark=True)}</div>
      <div class="wm" style="font-size:20px;color:#fff;letter-spacing:1px;margin-bottom:8px">SMART <span style="color:var(--cyan)">HSE</span> CHILE</div>
      <p>Gestión HSE para todas las áreas laborales de Chile · Cumplimiento DS.44</p>
      <p style="margin-top:12px"><a href="mailto:contacto@smarthse.cl">contacto@smarthse.cl</a> &nbsp;·&nbsp; <a href="https://smarthse.cl">smarthse.cl</a></p>
      <div class="ftr-sep">© 2025 Smart HSE Chile · Todos los derechos reservados.</div>
    </div>
    """, unsafe_allow_html=True)
# ════════════════════════════════════════════════════════════
# VISTA LOGIN
# ════════════════════════════════════════════════════════════
def login():
    st.markdown("<br><br><br>",unsafe_allow_html=True)
    _,col,_ = st.columns([1,1.2,1])
    with col:
        st.markdown("<div style='text-align:center;margin-bottom:28px'><div style='font-family:Montserrat,sans-serif;font-weight:900;font-size:2.6rem;color:#002B49'>SMART HSE</div><div style='display:inline-block;background:#55B4B0;color:white;font-size:10px;font-weight:700;padding:2px 10px;border-radius:4px;letter-spacing:2px;margin-top:4px'>CHILE</div><p style='color:#64748b;margin-top:14px;font-weight:300;font-size:14px'>Consola de Gestión Operativa · DS.44</p></div>",unsafe_allow_html=True)
        with st.form("login"):
            st.text_input("Correo electrónico",placeholder="correo@empresa.cl")
            pw=st.text_input("Contraseña",type="password")
            ca,cb=st.columns(2)
            volver=ca.form_submit_button("← Volver",use_container_width=True)
            entrar=cb.form_submit_button("Ingresar →",type="primary",use_container_width=True)
        if entrar:
            if pw==APP_PW: st.session_state.update({"auth":True,"vista":"consola"}); st.rerun()
            else: st.error("Contraseña incorrecta.")
        if volver: st.session_state["vista"]="landing"; st.rerun()

# ════════════════════════════════════════════════════════════
# VISTA CONSOLA
# ════════════════════════════════════════════════════════════
def consola():
    st.markdown("<style>section[data-testid='stSidebar']{display:flex!important}.block-container{padding:2rem!important;max-width:100%!important}</style>",unsafe_allow_html=True)
    with st.sidebar:
        st.markdown("<div style='text-align:center;padding:14px 0 6px'><div style='font-family:Montserrat,sans-serif;font-weight:900;font-size:20px;color:#002B49'>SMART HSE</div><div style='display:inline-block;background:#55B4B0;color:white;font-size:9px;font-weight:700;padding:1px 8px;border-radius:3px;letter-spacing:2px'>CHILE</div></div>",unsafe_allow_html=True)
        st.markdown("---")
        st.subheader("Selección de Faena")
        cliente=st.selectbox("Cliente:",list(ESTRUCTURA.keys()))
        faena=st.selectbox("Faena:",ESTRUCTURA[cliente]["faenas"])
        st.session_state["ruta"]=ESTRUCTURA[cliente]["ruta"]
        st.success(f"📁 {st.session_state['ruta']}")
        st.markdown("---")
        if st.button("🔒 Cerrar Sesión",use_container_width=True):
            st.session_state.update({"auth":False,"vista":"landing"}); st.rerun()

    st.title(f"Gestión Operativa: {faena}")
    st.caption(f"Smart HSE Chile · {datetime.today().strftime('%d/%m/%Y')} · Asesora Senior")

    t1,t2,t3,t4,t5,t6 = st.tabs([
        "⚡ Acción Inmediata",
        "📥 Motor Documental",
        "📊 GAP Analysis",
        "📄 Cartas N/A",
        "📅 Agenda",
        "🚨 Incidentes",
    ])

    # ── TAB 1: Dashboard ────────────────────────────────────
    with t1:
        c1,c2,c3=st.columns(3)
        c1.markdown("<div class='kpi'><div class='val'>12</div><div class='lbl'>Documentos RESSO</div></div>",unsafe_allow_html=True)
        c2.markdown("<div class='kpi'><div class='val'>45</div><div class='lbl'>Registros FYS</div></div>",unsafe_allow_html=True)
        c3.markdown("<div class='kpi'><div class='val'>8</div><div class='lbl'>Ítems N/A detectados</div></div>",unsafe_allow_html=True)
        st.divider()
        st.markdown("#### ⚡ Foco del Día — Top 3")
        acts=sorted([a for a in ACTIVIDADES if a["cliente"]==cliente],key=lambda x:x["dias"])[:3]
        if not acts: st.info("Sin tareas urgentes para este cliente.")
        for i,a in enumerate(acts,1):
            px=(datetime.today()+timedelta(days=a["dias"])).strftime("%d/%m/%Y")
            cl="ar" if a["dias"]<=1 else "aa" if a["dias"]<=7 else "ag"
            ic="🔴" if a["dias"]<=1 else "⚠️" if a["dias"]<=7 else "✅"
            st.markdown(f'<div class="{cl}"><b>{i}. {ic} {a["actividad"]}</b><br><span style="font-size:.83rem">📅 Próxima: {px} · Módulo: {a["modulo"]}</span></div>',unsafe_allow_html=True)
        reqs=REQUISITOS.get(cliente,[])
        if reqs:
            st.divider(); st.markdown("#### 📋 Requisitos Formales del Cliente")
            for i,r in enumerate(reqs,1): st.markdown(f"**{i}.** {r}")

    # ── TAB 2: Motor Documental ─────────────────────────────
    with t2:
        st.subheader("Motor Documental — Ingreso y Clasificación")
        st.caption("Sube el documento, selecciona contrato y flujo. El sistema genera el nombre correcto y te lo entrega listo para archivar.")
        archivo=st.file_uploader("Selecciona el documento",type=["pdf","docx","xlsx","csv"])
        if archivo:
            ext=archivo.name.rsplit(".",1)[-1].lower()
            if ext in ["xlsx","csv"]:
                try:
                    df_prev=pd.read_excel(archivo) if ext=="xlsx" else pd.read_csv(archivo)
                    st.dataframe(df_prev,use_container_width=True)
                    st.caption(f"{len(df_prev)} filas · {len(df_prev.columns)} columnas")
                except Exception as e: st.warning(f"Vista previa no disponible: {e}")
            else:
                st.info(f"📎 **{archivo.name}** · {round(archivo.size/1024,1)} KB")
            st.divider()
            ca,cb=st.columns(2)
            contrato_m=ca.selectbox("Contrato",CONTRATOS)
            flujo_m=cb.selectbox("Flujo documental",FLUJOS)
            codigo="FYS" if flujo_m=="FYS Diario" else "RESSO"
            nombre_nuevo=f"{date.today().isoformat()}_{codigo}_{contrato_m}.{ext}"
            st.markdown(f"**Nombre generado:** `{nombre_nuevo}`")
            st.markdown(f"**Carpeta destino:** `{'05_Salida_FYS/'+contrato_m if flujo_m=='FYS Diario' else '04_Salida_RESSO'}/`")
            archivo.seek(0)
            st.download_button(
                label="⬇️ Descargar archivo renombrado",
                data=archivo.read(),
                file_name=nombre_nuevo,
                mime=archivo.type,
                type="primary",
                use_container_width=True
            )
        else:
            st.info("Sube un documento para comenzar el proceso de clasificación.")

    # ── TAB 3: GAP Analysis ─────────────────────────────────
    with t3:
        st.subheader("Base Documental y GAP Analysis")
        st.caption("Sube tu matriz Excel o CSV con columnas de estado/brecha. El sistema detecta automáticamente los ítems con incumplimiento.")
        archivos_gap=st.file_uploader("Cargar matrices GAP/FYS",type=["xlsx","xls","csv"],accept_multiple_files=True)
        if archivos_gap:
            for f in archivos_gap:
                ext_g=f.name.rsplit(".",1)[-1].lower()
                st.markdown(f"---\n#### 📋 {f.name}")
                try:
                    df=pd.read_excel(f) if ext_g in ["xlsx","xls"] else pd.read_csv(f)
                    df=df.dropna(how="all")
                    total=len(df)
                    cm1,cm2,cm3=st.columns(3)
                    cm1.metric("Total ítems",total)
                    col_est=next((c for c in df.columns if any(k in str(c).upper() for k in ["ESTADO","BRECHA","CUMPLIMIENTO","STATUS"])),None)
                    if col_est:
                        brechas=df[col_est].astype(str).str.upper().str.contains("NO|BRECHA|INCUMPLE|PENDIENTE",na=False).sum()
                        cumplidos=df[col_est].astype(str).str.upper().str.contains("SI|CUMPLE|OK|COMPLETO",na=False).sum()
                        cm2.metric("Brechas",int(brechas),delta=f"-{int(brechas)}",delta_color="inverse")
                        cm3.metric("Cumplidos",int(cumplidos),delta=f"+{int(cumplidos)}")
                    else:
                        cm2.metric("Columna estado","No detectada")
                        cm3.metric("Filas con datos",total)
                    st.dataframe(df.astype(str),use_container_width=True)
                except Exception as e: st.error(f"Error leyendo {f.name}: {e}")
        else:
            st.info("Sube archivos Excel/CSV con 'GAP', 'FYS' o 'RESSO' en el nombre para activar el análisis de brechas normativas.")

    # ── TAB 4: Cartas N/A ───────────────────────────────────
    with t4:
        st.subheader("Generador de Cartas de No Aplicabilidad")
        st.caption("Carga masiva desde Excel → genera una carta Word por cada ítem. General: sirve para cualquier empresa y cualquier cliente o mandante (no exclusivo de Codelco/RESSO).")

        cpl1,cpl2 = st.columns([3,2])
        with cpl1:
            ciudad = st.text_input("Ciudad del encabezado", value="Antofagasta")
        with cpl2:
            st.download_button(
                "⬇️ Descargar plantilla Excel",
                data=plantilla_excel_bytes(),
                file_name="Plantilla_Cartas_NA.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )

        archivo = st.file_uploader(
            "Sube el Excel con los ítems N/A (una fila = una carta)",
            type=["xlsx","xls"],
        )

        st.markdown(
            '<div class="nota">Columnas reconocidas (tolerante a tildes y typos): '
            '<strong>Empresa, Fecha, Señores, Ref., Nombre del contrato, N° contrato, '
            'Declaración de No Aplica, Nombre del cliente, Responsable, Cargo</strong>. '
            '<em>Punto de RESSO</em> es opcional.</div>',
            unsafe_allow_html=True,
        )

        if archivo is not None:
            try:
                xl = pd.ExcelFile(archivo)
                hoja = st.selectbox("Hoja a procesar", xl.sheet_names,
                                    index=(xl.sheet_names.index("Carta de N A")
                                           if "Carta de N A" in xl.sheet_names else 0))
                df = mapear_columnas(pd.read_excel(xl, sheet_name=hoja)).dropna(how="all")
                df = df[df.apply(lambda r: bool(_cval(r,"declaracion") or _cval(r,"ref")), axis=1)]

                if df.empty:
                    st.warning("No se encontraron filas con 'Declaración de No Aplica' o 'Ref.'. Revisa la planilla.")
                else:
                    st.success(f"{len(df)} ítem(s) detectado(s). Vista previa:")
                    cols_prev = [c for c in ["empresa","ref","nombre_cliente","declaracion"] if c in df.columns]
                    st.dataframe(df[cols_prev] if cols_prev else df, use_container_width=True, height=220)

                    if st.button("📄 Generar cartas (.docx) y empaquetar ZIP",
                                 type="primary", use_container_width=True):
                        zbuf = io.BytesIO()
                        with zipfile.ZipFile(zbuf, "w", zipfile.ZIP_DEFLATED) as zf:
                            for idx, row in df.reset_index(drop=True).iterrows():
                                dbuf = io.BytesIO()
                                construir_carta_doc(row, ciudad).save(dbuf)
                                zf.writestr(nombre_archivo_carta(row, idx), dbuf.getvalue())
                        st.download_button(
                            label=f"⬇️ Descargar {len(df)} carta(s) (.zip)",
                            data=zbuf.getvalue(),
                            file_name=f"Cartas_No_Aplica_{date.today().isoformat()}.zip",
                            mime="application/zip",
                            use_container_width=True,
                        )
                        st.markdown('<div class="nota">⚠️ <strong>Nota:</strong> Inserte logos corporativos y firma en cada Word antes de la presentación oficial.</div>',unsafe_allow_html=True)
            except Exception as e:
                st.error(f"No se pudo procesar el Excel: {e}")

    # ── TAB 6: Registro de Incidentes ──────────────────────
    with t6:
        st.subheader("🚨 Registro de Incidentes y Eventos HSE")
        st.caption("Registra accidentes, incidentes de alto potencial, casi accidentes y condiciones inseguras. Genera folio automático y descarga el registro en Excel.")

        TIPOS_EVENTO = [
            "Accidente con Tiempo Perdido (ATP)",
            "Accidente sin Tiempo Perdido (ASTP)",
            "Incidente de Alto Potencial (IAP)",
            "Casi Accidente (CA)",
            "Condición Insegura (CI)",
            "Acto Inseguro (AI)",
            "Enfermedad Profesional (EP)",
        ]
        SEVERIDADES = ["Leve", "Moderado", "Grave", "Fatal"]
        PARTES_CUERPO = [
            "Cabeza","Cuello","Hombro derecho","Hombro izquierdo",
            "Brazo derecho","Brazo izquierdo","Mano derecha","Mano izquierda",
            "Tórax / Espalda","Abdomen","Cadera","Pierna derecha","Pierna izquierda",
            "Rodilla derecha","Rodilla izquierda","Pie derecho","Pie izquierdo",
            "Ojos","Sin lesión física",
        ]

        with st.form("form_incidente", clear_on_submit=True):
            st.markdown("#### 📋 Datos del Evento")
            ri1, ri2, ri3 = st.columns(3)
            fecha_inc  = ri1.date_input("Fecha del evento", value=date.today())
            hora_inc   = ri2.time_input("Hora del evento", value=datetime.now().time())
            tipo_ev    = ri3.selectbox("Tipo de evento", TIPOS_EVENTO)

            ri4, ri5 = st.columns(2)
            contrato_inc = ri4.selectbox("Contrato", CONTRATOS)
            lugar_inc    = ri5.text_input("Lugar / Área específica", placeholder="Ej: Patio norte, Nivel -120, Sector chancado")

            desc_inc = st.text_area("Descripción del evento", height=100,
                placeholder="Describa qué ocurrió, cómo ocurrió y las condiciones del entorno al momento del evento.")

            st.markdown("#### 👷 Personas Involucradas")
            ri6, ri7 = st.columns(2)
            trabajador = ri6.text_input("Nombre trabajador(es)", placeholder="Nombre completo o 'Sin lesionados'")
            testigos   = ri7.text_input("Testigo(s)", placeholder="Nombre(s) o 'Sin testigos'")

            ri8, ri9 = st.columns(2)
            parte_cuerpo = ri8.selectbox("Parte del cuerpo afectada", PARTES_CUERPO)
            severidad    = ri9.selectbox("Severidad", SEVERIDADES)

            st.markdown("#### ⚡ Respuesta Inmediata")
            acciones_imm = st.text_area("Acciones inmediatas tomadas", height=80,
                placeholder="Ej: Se aisló el área, se prestó primeros auxilios, se trasladó al policlínico...")
            requiere_inv = st.checkbox("¿Requiere investigación formal?", value=(tipo_ev in [
                "Accidente con Tiempo Perdido (ATP)",
                "Incidente de Alto Potencial (IAP)",
                "Accidente sin Tiempo Perdido (ASTP)",
            ]))

            registrar = st.form_submit_button("🚨 Registrar Evento", type="primary", use_container_width=True)

        if registrar:
            n = len(st.session_state["incidentes"]) + 1
            folio = f"INC-{fecha_inc.strftime('%Y%m%d')}-{n:03d}"
            nuevo = {
                "Folio":           folio,
                "Fecha":           fecha_inc.isoformat(),
                "Hora":            hora_inc.strftime("%H:%M"),
                "Tipo de Evento":  tipo_ev,
                "Contrato":        contrato_inc,
                "Lugar":           lugar_inc,
                "Descripción":     desc_inc,
                "Trabajador(es)":  trabajador,
                "Testigos":        testigos,
                "Parte Cuerpo":    parte_cuerpo,
                "Severidad":       severidad,
                "Acciones Inmediatas": acciones_imm,
                "Requiere Investigación": "Sí" if requiere_inv else "No",
                "Registrado":      datetime.now().strftime("%Y-%m-%d %H:%M"),
                "Faena":           faena,
            }
            st.session_state["incidentes"].append(nuevo)
            if severidad in ["Grave","Fatal"] or "Alto Potencial" in tipo_ev:
                st.error(f"⚠️ **EVENTO CRÍTICO registrado** — Folio: `{folio}`. Notifique de inmediato a la supervisión y al cliente.")
            else:
                st.success(f"✅ Evento registrado con folio **{folio}**.")

        # ── Historial de incidentes ──
        st.divider()
        st.markdown("#### 📂 Historial de Eventos Registrados en Sesión")
        inc_list = st.session_state["incidentes"]

        if not inc_list:
            st.info("Aún no hay eventos registrados en esta sesión.")
        else:
            df_inc = pd.DataFrame(inc_list)

            # KPIs rápidos
            ki1,ki2,ki3,ki4 = st.columns(4)
            ki1.metric("Total eventos", len(df_inc))
            ki2.metric("ATP / ASTP", int(df_inc["Tipo de Evento"].str.contains("Accidente").sum()))
            ki3.metric("Alto Potencial", int(df_inc["Tipo de Evento"].str.contains("Alto Potencial").sum()))
            ki4.metric("Requieren investigación", int((df_inc["Requiere Investigación"]=="Sí").sum()))

            st.dataframe(df_inc, use_container_width=True, hide_index=True)

            # ── Descarga Excel ──
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
                df_inc.to_excel(writer, index=False, sheet_name="Registro Incidentes")
                ws = writer.sheets["Registro Incidentes"]
                # Ancho de columnas automático
                for col in ws.columns:
                    max_len = max(len(str(cell.value or "")) for cell in col)
                    ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)
            buffer.seek(0)

            st.download_button(
                label="⬇️ Descargar Registro Excel",
                data=buffer,
                file_name=f"Registro_Incidentes_{faena}_{date.today().isoformat()}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                type="primary",
                use_container_width=True,
            )

            if st.button("🗑️ Limpiar registro de sesión", use_container_width=True):
                st.session_state["incidentes"] = []
                st.rerun()

    # ── TAB 5: Agenda ───────────────────────────────────────
    with t5:
        st.markdown("#### 📅 Agenda de Actividades Recurrentes")
        df_ag=pd.DataFrame(ACTIVIDADES)
        df_ag.columns=["Actividad","Frecuencia","Días","Módulo","Cliente"]
        st.dataframe(df_ag,use_container_width=True,hide_index=True)

        st.divider(); st.subheader("🚀 Onboarding — Datos del Contrato")
        st.info(f"Estructura activa: **{st.session_state['ruta']}**")
        with st.form("onboarding"):
            oc=st.text_input("N° Contrato:",value=st.session_state["contrato"])
            oa=st.text_input("Actividad / Macro Proceso:",value=st.session_state["actividad"])
            ol=st.text_input("Lugar / Faena:",value=st.session_state["lugar"])
            if st.form_submit_button("💾 Guardar",type="primary"):
                st.session_state.update({"contrato":oc,"actividad":oa,"lugar":ol})
                st.success("✅ Guardado en sesión.")

# ════════════════════════════════════════════════════════════
# ROUTER
# ════════════════════════════════════════════════════════════
v=st.session_state["vista"]
if   v=="landing": landing()
elif v=="login":   login()
elif v=="consola":
    if st.session_state["auth"]: consola()
    else: st.session_state["vista"]="login"; st.rerun()
