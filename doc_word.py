"""Conversor mínimo HTML → .docx para los documentos generados del FUF.

Los documentos generables (Política, Organigrama, Procedimiento…) se producen como HTML por
`catalogo_documentos_ds44._documento_html`, con un subconjunto acotado de etiquetas: h1/h2, p, b/strong,
ul/li y table/tr/td (meta y firmas). Este módulo mapea ese subconjunto a un `python-docx.Document`, de
modo que el mismo documento pueda descargarse como Word real, sin depender de una librería externa de
conversión. No pretende ser un motor HTML completo: cubre lo que emite nuestro propio envoltorio.
"""
import io
import re
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


_AZUL = RGBColor(0x00, 0x6a, 0x9b)


class _HTMLaDocx(HTMLParser):
    """Recorre el HTML y va construyendo el documento Word."""

    def __init__(self, doc):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.buffer = ''            # texto acumulado del bloque actual
        self.modo = None            # 'h1'|'h2'|'p'|'li'|'td'
        self.bold = False
        self.en_tabla = False
        self.fila = None            # lista de celdas de la fila actual
        self.tablas = []            # filas acumuladas de la tabla actual

    # ── helpers ──
    def _flush_inline(self):
        return re.sub(r'\s+', ' ', self.buffer).strip()

    def _add_parrafo(self, texto, estilo=None):
        if not texto:
            return
        p = self.doc.add_paragraph(style=estilo)
        run = p.add_run(texto)
        if self.modo == 'h1':
            run.bold = True; run.font.size = Pt(15); run.font.color.rgb = _AZUL
        elif self.modo == 'h2':
            run.bold = True; run.font.size = Pt(12); run.font.color.rgb = _AZUL
        return p

    # ── eventos ──
    def handle_starttag(self, tag, attrs):
        if tag in ('h1', 'h2', 'p', 'li'):
            self.buffer = ''; self.modo = tag
        elif tag in ('b', 'strong'):
            self.bold = True
        elif tag == 'br':
            self.buffer += '\n'
        elif tag == 'table':
            self.en_tabla = True; self.tablas = []
        elif tag == 'tr' and self.en_tabla:
            self.fila = []
        elif tag in ('td', 'th') and self.en_tabla:
            self.buffer = ''; self.modo = 'td'

    def handle_data(self, data):
        if self.modo:
            self.buffer += data

    def handle_endtag(self, tag):
        if tag in ('b', 'strong'):
            self.bold = False
            return
        texto = self._flush_inline()
        if tag == 'h1':
            self._add_parrafo(texto); self.modo = None
        elif tag == 'h2':
            self._add_parrafo(texto); self.modo = None
        elif tag == 'p':
            self._add_parrafo(texto); self.modo = None
        elif tag == 'li':
            self._add_parrafo(texto, estilo='List Bullet'); self.modo = None
        elif tag in ('td', 'th') and self.en_tabla:
            if self.fila is not None:
                self.fila.append(texto)
            self.modo = None
        elif tag == 'tr' and self.en_tabla:
            if self.fila:
                self.tablas.append(self.fila)
            self.fila = None
        elif tag == 'table':
            self._volcar_tabla()
            self.en_tabla = False; self.tablas = []


    def _volcar_tabla(self):
        filas = [f for f in self.tablas if any(c.strip() for c in f)]
        if not filas:
            return
        ncols = max(len(f) for f in filas)
        t = self.doc.add_table(rows=0, cols=ncols)
        t.style = 'Table Grid'
        for f in filas:
            celdas = t.add_row().cells
            for i in range(ncols):
                celdas[i].text = f[i] if i < len(f) else ''
        self.doc.add_paragraph('')


def html_a_docx(html, titulo=None):
    """Convierte el HTML de un documento generado a bytes .docx."""
    doc = Document()
    if titulo:
        h = doc.add_paragraph()
        run = h.add_run(titulo)
        run.bold = True; run.font.size = Pt(16); run.font.color.rgb = _AZUL
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Solo el <body>: descarta <head>/<style> para no volcar CSS como texto.
    m = re.search(r'<body[^>]*>(.*)</body>', html, re.S | re.I)
    cuerpo = m.group(1) if m else html
    cuerpo = re.sub(r'<style.*?</style>', '', cuerpo, flags=re.S | re.I)
    cuerpo = re.sub(r'<button.*?</button>', '', cuerpo, flags=re.S | re.I)
    _HTMLaDocx(doc).feed(cuerpo)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
