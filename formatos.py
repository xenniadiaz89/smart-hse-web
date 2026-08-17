"""Formatos Word descargables por ítem FUF.

Cuando un ítem del FUF se declara "No Cumple" (la empresa no tiene la documentación), la app ofrece
descargar el **formato ya hecho** correspondiente, como Word rellenable y con el logo de Smart HSE en el
encabezado. Los archivos base viven en `formatos_smarthse/` (curados desde SMART_HSE_DOCUMENTACION); el
logo se incrusta al vuelo al servir, dejando el original intacto.
"""
import io
import os

_DIR = os.path.join(os.path.dirname(__file__), 'formatos_smarthse')
_LOGO = os.path.join(os.path.dirname(__file__), 'assets', 'logo_smarthse.png')
DOCX_MIME = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

# Ítem FUF → archivo de formato (Word) + nombre legible para la descarga.
FORMATO_POR_ITEM = {
    1:  ('Politica_SST.docx', 'Política de Seguridad y Salud en el Trabajo'),
    8:  ('Programa_Anual_SSO.docx', 'Programa Anual de Trabajo Preventivo'),
    9:  ('Programa_Anual_SSO.docx', 'Programa Anual de Trabajo Preventivo'),
    10: ('Programa_Anual_SSO.docx', 'Programa Anual de Trabajo Preventivo'),
    11: ('Programa_Anual_SSO.docx', 'Programa Anual de Trabajo Preventivo'),
    21: ('IRL_DS44.docx', 'Información de Riesgos Laborales (IRL)'),
    22: ('IRL_DS44.docx', 'Información de Riesgos Laborales (IRL)'),
    27: ('Plan_Emergencia.docx', 'Plan de Gestión de Emergencias'),
    30: ('Acta_Constitucion_CPHS.docx', 'Acta de Constitución del CPHS'),
    32: ('Acta_Constitucion_CPHS.docx', 'Acta de Constitución del CPHS'),
    35: ('Formato_Acta_Reunion_CPHS.docx', 'Formato de Acta de Reunión CPHS'),
    45: ('Designacion_DPR.docx', 'Designación del Encargado / DPR'),
    49: ('Reglamento_Interno_RIOHS.docx', 'Reglamento Interno de Higiene y Seguridad (RIOHS)'),
    50: ('Reglamento_Interno_RIOHS.docx', 'Reglamento Interno de Higiene y Seguridad (RIOHS)'),
    51: ('Reglamento_Interno_RIOHS.docx', 'Reglamento Interno de Higiene y Seguridad (RIOHS)'),
    52: ('Reglamento_Interno_RIOHS.docx', 'Reglamento Interno de Higiene y Seguridad (RIOHS)'),
}


def formato_de(n):
    """(ruta, nombre_legible) del formato Word del ítem FUF n, o None si no hay o falta el archivo."""
    try:
        n = int(n)
    except (TypeError, ValueError):
        return None
    par = FORMATO_POR_ITEM.get(n)
    if not par:
        return None
    ruta = os.path.join(_DIR, par[0])
    return (ruta, par[1]) if os.path.exists(ruta) else None


def tiene_formato(n):
    return formato_de(n) is not None


def brandear_docx(ruta):
    """Devuelve los bytes del .docx con el logo de Smart HSE incrustado en el encabezado.
    Best-effort: si python-docx o el logo faltan, devuelve el archivo original sin modificar."""
    try:
        import docx
        from docx.shared import Cm
        d = docx.Document(ruta)
        if os.path.exists(_LOGO):
            sec = d.sections[0]
            hdr = sec.header
            par = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
            # Evita duplicar el logo si el documento ya se braendó antes.
            if not par.runs:
                par.add_run().add_picture(_LOGO, height=Cm(1.1))
        buf = io.BytesIO()
        d.save(buf)
        return buf.getvalue()
    except Exception:
        with open(ruta, 'rb') as fh:
            return fh.read()
