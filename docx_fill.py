"""Motor de relleno de plantillas .docx (Word) preservando el formato original.

La usuaria pidió que los documentos grandes (Programa/SGSST V8.2, RIOHS v4) se descarguen con el
formato Word idéntico al original. En vez de reconstruirlos en HTML, se rellena el .docx real:
- **Etiquetas de tabla**: donde el formato trae una tabla "Etiqueta | (vacío)" (Antecedentes de la
  Empresa), se escribe el valor en la celda contigua manteniendo su estilo.
- **Tokens** `{{clave}}`: reemplazo directo en párrafos y celdas, por si una plantilla los usa.

Las plantillas viven EN el repo (`plantillas_docx/`) para que funcione también en Render (el folder
DS44/ está fuera del repo). Devuelve los bytes del .docx para descargar/persistir.
"""
import io
import os
import re
import unicodedata

PLANTILLAS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'plantillas_docx')


def _norm(s):
    s = (s or '').lower().strip().rstrip(':').strip()
    return ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')


def _set_cell_text(cell, valor):
    """Escribe `valor` en la celda conservando el estilo del primer run existente."""
    valor = '' if valor is None else str(valor)
    par = cell.paragraphs[0]
    if par.runs:
        par.runs[0].text = valor
        for r in par.runs[1:]:
            r.text = ''
    else:
        par.add_run(valor)


def _replace_tokens_paragraph(par, valores, reemplazos=None):
    """Reemplaza {{clave}} y patrones regex en un párrafo. Une los runs para no partir el match."""
    texto = par.text
    nuevo = texto
    for k, v in (valores or {}).items():
        nuevo = nuevo.replace('{{' + k + '}}', '' if v is None else str(v))
    for pat, val in (reemplazos or []):
        nuevo = re.sub(pat, '' if val is None else str(val), nuevo)
    if nuevo != texto:
        if par.runs:
            par.runs[0].text = nuevo
            for r in par.runs[1:]:
                r.text = ''
        else:
            par.add_run(nuevo)


def _tabla_por_cabecera(doc, cabecera):
    """La PRIMERA tabla cuya fila 1 coincide con `cabecera`, o None.

    Por cabecera y no por índice: la plantilla del Programa tiene 16 tablas y cualquier reedición
    del Word en el futuro desplazaría los índices sin avisar. «Primera» importa: en el V8.2 hay dos
    tablas ACTIVIDAD/PERIODICIDAD/RESPONSABLE — la primera es el programa de la empresa y la
    segunda el de proyectos en ejecución. Las medidas de la MIPER son de empresa.
    """
    buscada = [_norm(c) for c in cabecera]
    for tab in doc.tables:
        if not tab.rows:
            continue
        celdas = [_norm(c.text) for c in tab.rows[0].cells]
        if len(celdas) >= len(buscada) and celdas[:len(buscada)] == buscada:
            return tab
    return None


def _agregar_fila(tab, textos, modelo=None):
    """Añade una fila clonando `modelo` (o la última) para heredar el formato del documento real."""
    import copy
    fuente = modelo if modelo is not None else tab.rows[-1]
    tr = copy.deepcopy(fuente._tr)
    tab._tbl.append(tr)
    fila = tab.rows[-1]
    for i, celda in enumerate(fila.cells):
        _set_cell_text(celda, textos[i] if i < len(textos) else '')
    return fila


def _fila_de_seccion(tab):
    """Una fila 'separador de sección' de la tabla (las que repiten el mismo texto en toda la
    fila, como «POLITICA, LIDERAZGO Y COMPROMISO»), para clonar su formato destacado."""
    for row in tab.rows[1:]:
        textos = [(c.text or '').strip() for c in row.cells]
        if textos[0] and len(set(textos)) == 1:
            return row
    return None


def agregar_filas(doc, cabecera, filas, titulo_seccion=None):
    """Añade `filas` al final de la tabla identificada por `cabecera`. Devuelve cuántas escribió.

    Best-effort: si la tabla no está en la plantilla, no hace nada y devuelve 0 — el documento se
    genera igual, que es lo que importa.
    """
    tab = _tabla_por_cabecera(doc, cabecera)
    if tab is None or not filas:
        return 0
    modelo = tab.rows[-1]
    if titulo_seccion:
        seccion = _fila_de_seccion(tab)
        if seccion is not None:
            _agregar_fila(tab, [titulo_seccion] * len(tab.columns), modelo=seccion)
    for f in filas:
        _agregar_fila(tab, f, modelo=modelo)
    return len(filas)


def rellenar(nombre_plantilla, valores=None, etiquetas=None, reemplazos=None, tablas=None):
    """Abre `plantillas_docx/<nombre_plantilla>`, rellena y devuelve (bytes, filename).
    - `valores`: dict para reemplazo de {{tokens}} (párrafos y tablas).
    - `etiquetas`: dict {texto_de_etiqueta_normalizado: valor} → rellena la celda contigua a la
      etiqueta en cualquier tabla (para tablas "Etiqueta | valor").
    - `reemplazos`: lista de (patrón_regex, valor) → reemplazo literal (ej. placeholders XXXX).
    - `tablas`: lista de {'cabecera': [...], 'filas': [[celda, ...], ...], 'titulo': str} → añade
      filas al final de la tabla que tenga esa cabecera (ej. las medidas de la MIPER).
    Best-effort: si python-docx o la plantilla faltan, devuelve (None, None).
    """
    valores = valores or {}
    etiquetas_norm = {_norm(k): v for k, v in (etiquetas or {}).items()}
    ruta = os.path.join(PLANTILLAS_DIR, nombre_plantilla)
    if not os.path.exists(ruta):
        return None, None
    try:
        import docx
    except ImportError:
        return None, None
    try:
        d = docx.Document(ruta)
        for par in d.paragraphs:
            _replace_tokens_paragraph(par, valores, reemplazos)
        for tab in d.tables:
            for row in tab.rows:
                cells = row.cells
                for i, cell in enumerate(cells):
                    for par in cell.paragraphs:
                        _replace_tokens_paragraph(par, valores, reemplazos)
                    # tabla etiqueta|valor: si esta celda es una etiqueta conocida, rellena la siguiente
                    if etiquetas_norm and i + 1 < len(cells):
                        key = _norm(cell.text)
                        if key in etiquetas_norm and not cells[i + 1].text.strip():
                            _set_cell_text(cells[i + 1], etiquetas_norm[key])
        for t in (tablas or []):
            agregar_filas(d, t.get('cabecera') or [], t.get('filas') or [], t.get('titulo'))
        buf = io.BytesIO()
        d.save(buf)
        return buf.getvalue(), nombre_plantilla
    except Exception:      # noqa: BLE001 — best-effort; el pipeline degrada a HTML si falla
        return None, None


# ── Documentos .docx disponibles: mapea tipo_doc → plantilla + cómo auto-llenar desde la empresa ──
def _etiquetas_empresa(empresa):
    e = empresa or {}
    return {
        'razon social': e.get('razon_social'),
        'rut': e.get('rut_empresa'),
        'nombre representante legal': e.get('representante') or '',
        'domicilio comercial': e.get('direccion') or '',
        'organismo administrador ley': e.get('mutual') or '',
        'numero adherente mutual de seguridad': e.get('n_adherente') or '',
    }


def _reemplazos_riohs(empresa, campos):
    # El RIOHS v4 usa 'XXXX…' como marcador del nombre de la empresa (50 apariciones).
    razon = (empresa or {}).get('razon_social') or 'LA EMPRESA'
    return [(r'X{4,}', razon)]


DOCX = {
    'programa_sgsst': {
        'plantilla': 'programa_sgsst_v8.2.docx',
        'nombre': 'Programa Anual y SGSST',
        'items_fuf': [8, 9, 10, 11, 12, 13, 20],
        'campos': [{'k': 'periodo', 'label': 'Período (año)', 'tipo': 'text'}],
        # Tabla de actividades donde se vuelcan las medidas de la MIPER (ítem 10 del FUF).
        'tabla_medidas': ['ACTIVIDAD', 'PERIODICIDAD', 'RESPONSABLE'],
        'etiquetas': _etiquetas_empresa,
        'tokens': lambda empresa, campos: {
            'razon_social': (empresa or {}).get('razon_social') or '',
            'periodo': (campos or {}).get('periodo') or '',
        },
    },
    'riohs_docx': {
        'plantilla': 'riohs_v4.docx',
        'nombre': 'Reglamento Interno (RIOHS) — formato Word oficial',
        'items_fuf': [49, 50, 51, 52],
        'campos': [],
        'reemplazos': _reemplazos_riohs,
    },
}


def tipos_para_item(n):
    """Documentos .docx generables para el ítem FUF n (para el catálogo/UI del FUF)."""
    try:
        n = int(n)
    except (TypeError, ValueError):
        return []
    return [{'tipo_doc': k, 'nombre': cfg['nombre'], 'campos': cfg.get('campos', []), 'formato': 'docx'}
            for k, cfg in DOCX.items() if n in cfg['items_fuf']]


def es_docx(tipo_doc):
    return tipo_doc in DOCX


def filas_medidas_miper(medidas):
    """Convierte las medidas de db.medidas_para_programa() en filas ACTIVIDAD|PERIODICIDAD|RESPONSABLE.

    La PERIODICIDAD sale en blanco a propósito: el DS 44 no fija plazos por nivel de riesgo, y
    escribir uno inventado en un documento oficial lo convertiría en un compromiso que la empresa
    no ha asumido. Lo completa el experto.
    """
    filas = []
    for m in (medidas or []):
        origen = ' · '.join(x for x in (m.get('proceso'), m.get('tarea'), m.get('puesto')) if x)
        riesgo = m.get('riesgo') or m.get('peligro') or ''
        actividad = m.get('medida') or ''
        if riesgo:
            actividad = f"[{riesgo}] {actividad}"
        if origen:
            actividad = f"{actividad}\n({origen})"
        filas.append([actividad, '', m.get('responsable') or ''])
    return filas


def generar_docx(tipo_doc, empresa, campos=None, medidas_miper=None):
    """Genera el .docx de `tipo_doc` auto-llenado desde la empresa. Devuelve (bytes, filename).

    `medidas_miper` (opcional) son las medidas de control de la matriz vigente: se vuelcan a la
    tabla de actividades del Programa, que es lo que exige el ítem 10 del FUF.
    """
    cfg = DOCX.get(tipo_doc)
    if not cfg:
        return None, None
    etiquetas = cfg['etiquetas'](empresa) if callable(cfg.get('etiquetas')) else cfg.get('etiquetas')
    tokens = cfg['tokens'](empresa, campos) if callable(cfg.get('tokens')) else cfg.get('tokens')
    reemplazos = cfg['reemplazos'](empresa, campos) if callable(cfg.get('reemplazos')) else cfg.get('reemplazos')
    tablas = None
    if medidas_miper and cfg.get('tabla_medidas'):
        tablas = [{'cabecera': cfg['tabla_medidas'],
                   'titulo': 'MEDIDAS PREVENTIVAS DERIVADAS DE LA MATRIZ DE RIESGOS (MIPER)',
                   'filas': filas_medidas_miper(medidas_miper)}]
    data, _ = rellenar(cfg['plantilla'], valores=tokens, etiquetas=etiquetas,
                       reemplazos=reemplazos, tablas=tablas)
    if data is None:
        return None, None
    emp = (empresa or {}).get('razon_social') or 'empresa'
    return data, f"{cfg['nombre']} - {emp}.docx"
